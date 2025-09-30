# backend/app/services/document_processor.py
"""
Simple document processor for AURA Voice AI
Handles file uploads and extracts text for AI context
"""

import os
import json
import logging
from typing import Dict, List, Optional
from dataclasses import dataclass
import hashlib
from datetime import datetime
import PyPDF2
import docx
import markdown

logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Simple document storage"""
    id: str
    filename: str
    content: str
    file_type: str
    upload_time: str
    size: int
    user_id: Optional[str] = None

class DocumentProcessor:
    def __init__(self, storage_path: str = "./document_storage"):
        """Initialize document processor with local storage"""
        self.storage_path = storage_path
        self.documents = {}  # In-memory cache
        
        # Create storage directory if not exists
        if not os.path.exists(storage_path):
            os.makedirs(storage_path)
            
        logger.info(f"Document processor initialized with storage at {storage_path}")
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from various file formats"""
        try:
            if file_type == 'txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
                    
            elif file_type == 'pdf':
                text = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text
                
            elif file_type in ['doc', 'docx']:
                doc = docx.Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
                
            elif file_type == 'md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    # Convert markdown to plain text (remove formatting)
                    return content  # Keep markdown for now
                    
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    async def process_upload(self, file_data: bytes, filename: str, user_id: Optional[str] = None) -> Document:
        """Process uploaded file and extract content"""
        try:
            # Generate unique document ID
            doc_id = hashlib.md5(f"{filename}{datetime.now().isoformat()}".encode()).hexdigest()[:12]
            
            # Determine file type
            file_extension = filename.split('.')[-1].lower()
            
            # Save file temporarily
            temp_path = os.path.join(self.storage_path, f"{doc_id}_{filename}")
            with open(temp_path, 'wb') as f:
                f.write(file_data)
            
            # Extract text content
            content = self.extract_text(temp_path, file_extension)
            
            # Create document object
            document = Document(
                id=doc_id,
                filename=filename,
                content=content,
                file_type=file_extension,
                upload_time=datetime.now().isoformat(),
                size=len(file_data),
                user_id=user_id
            )
            
            # Store in memory and save metadata
            self.documents[doc_id] = document
            self._save_metadata(document)
            
            logger.info(f"Processed document {filename} with ID {doc_id}")
            return document
            
        except Exception as e:
            logger.error(f"Error processing upload: {e}")
            raise
    
    def _save_metadata(self, document: Document):
        """Save document metadata to disk"""
        metadata_path = os.path.join(self.storage_path, f"{document.id}_metadata.json")
        with open(metadata_path, 'w') as f:
            json.dump({
                'id': document.id,
                'filename': document.filename,
                'file_type': document.file_type,
                'upload_time': document.upload_time,
                'size': document.size,
                'user_id': document.user_id
            }, f)
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get document by ID"""
        # Check memory cache first
        if doc_id in self.documents:
            return self.documents[doc_id]
        
        # Try to load from disk
        metadata_path = os.path.join(self.storage_path, f"{doc_id}_metadata.json")
        if os.path.exists(metadata_path):
            with open(metadata_path, 'r') as f:
                metadata = json.load(f)
            
            # Load content
            file_path = os.path.join(self.storage_path, f"{doc_id}_{metadata['filename']}")
            content = self.extract_text(file_path, metadata['file_type'])
            
            document = Document(
                id=metadata['id'],
                filename=metadata['filename'],
                content=content,
                file_type=metadata['file_type'],
                upload_time=metadata['upload_time'],
                size=metadata['size'],
                user_id=metadata.get('user_id')
            )
            
            self.documents[doc_id] = document
            return document
        
        return None
    
    def get_user_documents(self, user_id: str) -> List[Document]:
        """Get all documents for a user"""
        user_docs = []
        
        # Check all metadata files
        for filename in os.listdir(self.storage_path):
            if filename.endswith('_metadata.json'):
                with open(os.path.join(self.storage_path, filename), 'r') as f:
                    metadata = json.load(f)
                    if metadata.get('user_id') == user_id:
                        doc = self.get_document(metadata['id'])
                        if doc:
                            user_docs.append(doc)
        
        return user_docs
    
    def search_documents(self, query: str, user_id: Optional[str] = None) -> List[Dict]:
        """Simple text search in documents"""
        results = []
        query_lower = query.lower()
        
        # Get relevant documents
        if user_id:
            docs = self.get_user_documents(user_id)
        else:
            docs = list(self.documents.values())
        
        for doc in docs:
            if query_lower in doc.content.lower():
                # Find relevant snippet
                content_lower = doc.content.lower()
                index = content_lower.find(query_lower)
                start = max(0, index - 100)
                end = min(len(doc.content), index + 200)
                snippet = doc.content[start:end]
                
                results.append({
                    'document_id': doc.id,
                    'filename': doc.filename,
                    'snippet': snippet,
                    'relevance': 1.0  # Simple scoring
                })
        
        return results[:5]  # Return top 5 results
    
    def get_context_for_query(self, query: str, doc_id: Optional[str] = None, user_id: Optional[str] = None) -> str:
        """Get relevant context for AI query"""
        context = ""
        
        # If specific document requested
        if doc_id:
            doc = self.get_document(doc_id)
            if doc:
                context = f"Based on the document '{doc.filename}':\n\n{doc.content[:3000]}"
        
        # Otherwise search for relevant content
        elif user_id:
            results = self.search_documents(query, user_id)
            if results:
                context = "Relevant information from your documents:\n\n"
                for result in results[:3]:
                    context += f"From '{result['filename']}':\n{result['snippet']}\n\n"
        
        return context
    
    def clear_user_documents(self, user_id: str) -> bool:
        """Clear all documents for a user"""
        try:
            docs = self.get_user_documents(user_id)
            for doc in docs:
                # Remove files
                os.remove(os.path.join(self.storage_path, f"{doc.id}_{doc.filename}"))
                os.remove(os.path.join(self.storage_path, f"{doc.id}_metadata.json"))
                # Remove from cache
                if doc.id in self.documents:
                    del self.documents[doc.id]
            
            logger.info(f"Cleared {len(docs)} documents for user {user_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error clearing documents: {e}")
            return False